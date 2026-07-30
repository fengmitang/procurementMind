"""Create the V1.4 interface-document revision for Agent backend support."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parents[1] / "docs"


def set_metadata(document: Document, key: str, value: str) -> None:
    for row in document.tables[0].rows:
        if row.cells[0].text.strip() == key:
            row.cells[1].text = value
            return
    raise ValueError(f"Metadata key not found: {key}")


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    style: str | None = None,
) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    inserted = Paragraph(new_element, paragraph._parent)
    if style:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def clone_message_query_table(document: Document) -> Table:
    post_table = next(
        table
        for table in document.tables
        if len(table.rows) >= 4
        and table.rows[2].cells[1].text.strip()
        == "/api/v1/agent/conversations/{conversation_id}/messages"
        and table.rows[1].cells[1].text.strip() == "POST"
    )
    cloned_xml = deepcopy(post_table._tbl)
    post_table._tbl.addnext(cloned_xml)
    cloned = Table(cloned_xml, post_table._parent)
    cloned.rows[1].cells[1].text = "GET"
    cloned.rows[3].cells[1].text = "分页查询完整会话消息，供外部 Agent 重启、恢复或审计时重新读取。"
    return cloned


def update_interface_document() -> Path:
    source = next(DOCS.glob("*后端接口文档-v1.3.docx"))
    output = source.with_name(source.name.replace("v1.3", "v1.4"))
    document = Document(source)

    set_metadata(document, "文档版本", "V1.4")
    set_metadata(document, "文档状态", "Agent 后端支撑接口实现同步稿")
    set_metadata(document, "编制日期", "2026年7月30日")
    document.core_properties.title = "数据中心采购流程自动化 Agent 后端接口文档 V1.4"

    boundary = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("这些接口由 Agent 服务调用。")
    )
    boundary.text = (
        "这些接口由外部 Agent 服务调用。采购后端只负责会话、消息、结构化状态、"
        "Redis 短期缓存和 MySQL 快照，不负责大模型调用、提示词、意图识别、"
        "回复生成、工具选择或任务编排。实时状态存 Redis，完整消息与会话元数据"
        "存 MySQL，关键节点和结束状态写入 agent_session_state 快照。"
    )

    message_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "9.2 写入会话消息"
    )
    message_heading.text = "9.2 写入与查询会话消息"

    response_example = next(
        paragraph for paragraph in document.paragraphs if '"message_id": 9001' in paragraph.text
    )
    response_example.text = (
        '{\n  "message_id": 9001,\n'
        '  "created_at": "2026-07-29T10:00:00+08:00",\n'
        '  "duplicate": false\n}'
    )

    duplicate_rule = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("external_message_id 在同一会话内唯一")
    )
    query_note = insert_paragraph_after(
        duplicate_rule,
        "GET 同一路径按 page、page_size 分页返回消息，默认每页 50 条、最大 200 条，"
        "按 created_at 和 message_id 正序排列。外部 Agent 可据此在服务重启后恢复"
        "完整可见对话；接口不保存或返回模型内部思维过程。",
    )
    thought_rule = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "不保存模型内部思维过程。"
    )
    thought_rule._element.getparent().remove(thought_rule._element)
    query_note.style = document.styles["Normal"]

    restore_rule = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("每次读取或更新成功后刷新 TTL")
    )
    insert_paragraph_after(
        restore_rule,
        "Redis Key 不存在时，后端自动从 agent_session_state 的最新关键快照恢复并"
        "重新写入 Redis；MySQL 也不存在可恢复快照时返回 SESSION_EXPIRED。",
    )

    clone_message_query_table(document)
    document.save(output)
    return output


if __name__ == "__main__":
    print(update_interface_document())
