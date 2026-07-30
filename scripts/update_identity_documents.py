"""Create V1.2 identity-gateway revisions of the interface and technical documents."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parents[1] / "docs"


def copy_cell_style(source: _Cell, target: _Cell) -> None:
    target_tc_pr = target._tc.get_or_add_tcPr()
    for child in list(target_tc_pr):
        target_tc_pr.remove(child)
    source_tc_pr = source._tc.tcPr
    if source_tc_pr is not None:
        for child in source_tc_pr:
            target_tc_pr.append(deepcopy(child))


def append_row(table: Table, values: tuple[str, ...]) -> None:
    source_row = table.rows[-1]
    row = table.add_row()
    for index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
        copy_cell_style(source_row.cells[index], cell)
        cell.text = value


def set_metadata(table: Table, key: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == key:
            row.cells[1].text = value
            return
    raise ValueError(f"Metadata key not found: {key}")


def find_doc(version: str, name_codepoints: tuple[int, ...]) -> Path:
    marker = "".join(chr(codepoint) for codepoint in name_codepoints)
    return next(path for path in DOCS.glob(f"*{version}.docx") if marker in path.name)


def update_interface_document() -> Path:
    source = find_doc("v1.1", (21518, 31471, 25509, 21475))
    output = source.with_name(source.name.replace("v1.1", "v1.2"))
    document = Document(source)

    set_metadata(document.tables[0], "文档版本", "V1.2")
    set_metadata(document.tables[0], "文档状态", "身份网关安全约定修订稿")
    set_metadata(document.tables[0], "设计依据", "需求分析 V0.6、数据库设计 V1.3")

    identity_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "2.2 平台身份上下文"
    )
    body = document.paragraphs[identity_index + 1]
    example = document.paragraphs[identity_index + 2]
    note = document.paragraphs[identity_index + 3]

    body.text = (
        "本期采用“可信飞书适配层/内部网关 + 采购业务后端”模式。适配层先完成飞书"
        "请求验签和企业成员识别，再注入平台类型、平台用户唯一标识及网关签名。采购"
        "后端不接受客户端自报的 employee_id、姓名、手机号或角色，也不信任缺少有效"
        "网关签名的平台身份头；后端通过 employee_external_identity 将平台账号解析为"
        "本地 employee_id，并检查员工、外部身份、角色和楼宇关联是否有效。"
    )
    example.text = (
        "X-Platform-Type: FEISHU\n"
        "X-Platform-User-Id: ou_xxx\n"
        "X-Gateway-Timestamp: 1785312000\n"
        "X-Gateway-Nonce: 4f8f2c6e94e84b809aa92b8f21d49d7a\n"
        "X-Gateway-Signature: <HMAC-SHA256 十六进制摘要>\n"
        "X-Request-Id: 7e8b...（建议）"
    )
    note.text = (
        "网关签名原文依次为 HTTP 方法、URL 路径、平台类型、平台用户标识、时间戳和"
        "随机数，字段之间使用换行符连接；密钥只保存在适配层和业务后端的安全配置中。"
        "后端默认只接受 300 秒内的签名，并使用 Redis 记录随机数，在有效期内拒绝重复"
        "请求。开发环境允许 TEST_PLATFORM，生产环境禁止该测试平台类型。"
    )

    header_table = document.tables[2]
    append_row(
        header_table,
        ("X-Gateway-Timestamp", "是", "网关签名使用的 Unix 秒级时间戳"),
    )
    append_row(
        header_table,
        ("X-Gateway-Nonce", "是", "16 至 128 位随机数，有效期内不得重复"),
    )
    append_row(
        header_table,
        ("X-Gateway-Signature", "是", "网关使用共享密钥生成的 HMAC-SHA256 签名"),
    )

    document.core_properties.title = "数据中心采购流程自动化 Agent 后端接口文档 V1.2"
    document.save(output)
    return output


def insert_before(target: Paragraph, text: str, style: str) -> None:
    target.insert_paragraph_before(text=text, style=style)


def update_technical_document() -> Path:
    source = find_doc("v1.1", (25216, 26415, 36873, 22411))
    output = source.with_name(source.name.replace("v1.1", "v1.2"))
    document = Document(source)

    set_metadata(document.tables[0], "文档版本", "V1.2")
    set_metadata(
        document.tables[0],
        "对应文档",
        "需求分析 V0.6、数据库设计 V1.3、后端接口文档 V1.2",
    )

    next_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "9. 主要业务模块"
    )
    insert_before(next_heading, "8.6 身份网关与权限边界", "Heading 2")
    insert_before(
        next_heading,
        "飞书适配层或内部网关负责验证飞书请求和企业成员身份；采购后端只接受通过"
        "HMAC-SHA256 签名的平台身份上下文，不直接信任客户端传入的身份请求头。",
        "List Bullet",
    )
    insert_before(
        next_heading,
        "后端根据 employee_external_identity 解析员工，并校验员工状态、有效角色和"
        "楼宇范围。签名包含时间戳和随机数，Redis 用于短期防重放；生产环境不允许"
        "TEST_PLATFORM。",
        "List Bullet",
    )

    document.core_properties.title = "数据中心采购流程自动化 Agent 技术选型文档 V1.2"
    document.save(output)
    return output


if __name__ == "__main__":
    for result in (update_interface_document(), update_technical_document()):
        print(result)
