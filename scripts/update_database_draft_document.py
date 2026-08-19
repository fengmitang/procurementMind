"""Create database design V1.4 with nullable incomplete-draft fields."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parents[1] / "docs"


def set_metadata(document: Document, key: str, value: str) -> None:
    for row in document.tables[0].rows:
        if row.cells[0].text.strip() == key:
            row.cells[1].text = value
            return
    raise ValueError(f"Metadata key not found: {key}")


def main() -> None:
    database_marker = "".join(chr(codepoint) for codepoint in (25968, 25454, 24211))
    source = next(path for path in DOCS.glob("*v1.3.docx") if database_marker in path.name)
    output = source.with_name(source.name.replace("v1.3", "v1.4"))
    document = Document(source)

    set_metadata(document, "文档版本", "V1.4")
    set_metadata(document, "文档状态", "草稿字段约束修订稿")
    for key in ("对应需求版本", "需求版本", "设计依据"):
        try:
            set_metadata(
                document,
                key,
                "需求分析 V0.6、后端接口文档 V1.3",
            )
            break
        except ValueError:
            continue

    field_table = document.tables[11]
    draft_fields = {
        "device_profession": "设备专业或类别",
        "device_name": "设备名称",
        "quantity": "申请数量，非空时必须为正整数",
        "unit": "计量单位",
        "application_reason": "申请原因",
    }
    for row in field_table.rows[1:]:
        field_name = row.cells[0].text.strip()
        if field_name in draft_fields:
            row.cells[2].text = "草稿可空，提交前必填"
            row.cells[3].text = draft_fields[field_name]

    heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "14. purchase_request 采购申请表"
    )
    new_element = OxmlElement("w:p")
    heading._p.addnext(new_element)
    note = Paragraph(new_element, heading._parent)
    note.style = "Normal"
    note.add_run(
        "草稿创建接口只接收 building_id，因此设备专业、设备名称、数量、单位和申请原因"
        "在 DRAFT 或 REJECTED 编辑阶段允许为空；提交或重新提交楼长审核前，后端必须"
        "统一校验上述字段已填写完整。数据库 CHECK 约束仅在 quantity 非空时要求其为正整数。"
    )

    document.core_properties.title = "数据中心采购流程自动化 Agent 数据库设计文档 V1.4"
    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
