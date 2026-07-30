"""Synchronize interface and technical document references after database V1.4."""

from pathlib import Path

from docx import Document

DOCS = Path(__file__).resolve().parents[1] / "docs"


def find_document(version: str, marker_codepoints: tuple[int, ...]) -> Path:
    marker = "".join(chr(codepoint) for codepoint in marker_codepoints)
    return next(path for path in DOCS.glob(f"*{version}.docx") if marker in path.name)


def set_metadata(document: Document, key: str, value: str) -> None:
    for row in document.tables[0].rows:
        if row.cells[0].text.strip() == key:
            row.cells[1].text = value
            return
    raise ValueError(f"Metadata key not found: {key}")


def update_interface() -> Path:
    source = find_document("v1.2", (21518, 31471, 25509, 21475))
    output = source.with_name(source.name.replace("v1.2", "v1.3"))
    document = Document(source)
    set_metadata(document, "文档版本", "V1.3")
    set_metadata(document, "文档状态", "采购主流程实现同步稿")
    set_metadata(document, "设计依据", "需求分析 V0.6、数据库设计 V1.4")
    document.core_properties.title = "数据中心采购流程自动化 Agent 后端接口文档 V1.3"
    document.save(output)
    return output


def update_technical() -> Path:
    source = find_document("v1.2", (25216, 26415, 36873, 22411))
    output = source.with_name(source.name.replace("v1.2", "v1.3"))
    document = Document(source)
    set_metadata(document, "文档版本", "V1.3")
    set_metadata(
        document,
        "对应文档",
        "需求分析 V0.6、数据库设计 V1.4、后端接口文档 V1.3",
    )
    document.core_properties.title = "数据中心采购流程自动化 Agent 技术选型文档 V1.3"
    document.save(output)
    return output


if __name__ == "__main__":
    for path in (update_interface(), update_technical()):
        print(path)
