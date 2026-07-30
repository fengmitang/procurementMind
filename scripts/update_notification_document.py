"""Create the V1.5 interface-document revision for notification Outbox APIs."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCS = Path(__file__).resolve().parents[1] / "docs"


def set_metadata(document: Document, key: str, value: str) -> None:
    for row in document.tables[0].rows:
        if row.cells[0].text.strip() == key:
            row.cells[1].text = value
            return
    raise ValueError(f"Metadata key not found: {key}")


def add_endpoint_table_before(
    document: Document,
    target: Paragraph,
    *,
    method: str,
    path: str,
    purpose: str,
) -> Table:
    template = next(
        table
        for table in reversed(document.tables)
        if len(table.rows) == 5
        and table.rows[0].cells[0].text.strip() == "项目"
        and table.rows[1].cells[0].text.strip() == "方法"
    )
    cloned_xml = deepcopy(template._tbl)
    target._p.addprevious(cloned_xml)
    cloned = Table(cloned_xml, template._parent)
    cloned.rows[1].cells[1].text = method
    cloned.rows[2].cells[1].text = path
    cloned.rows[3].cells[1].text = purpose
    cloned.rows[4].cells[1].text = "需要平台身份上下文，仅 ADMIN"
    return cloned


def add_paragraph_before(
    target: Paragraph,
    text: str,
    style: str | None = None,
) -> Paragraph:
    return target.insert_paragraph_before(text=text, style=style)


def update_interface_document() -> Path:
    source = next(DOCS.glob("*后端接口文档-v1.4.docx"))
    output = source.with_name(source.name.replace("v1.4", "v1.5"))
    document = Document(source)

    set_metadata(document, "文档版本", "V1.5")
    set_metadata(document, "文档状态", "通知 Outbox 后端实现同步稿")
    set_metadata(document, "编制日期", "2026年7月30日")
    document.core_properties.title = "数据中心采购流程自动化 Agent 后端接口文档 V1.5"

    error_heading = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "10. 统一错误码"
    )
    add_paragraph_before(error_heading, "9.7 通知 Outbox 管理接口", "Heading 2")
    add_paragraph_before(
        error_heading,
        "通知记录由采购业务事务写入 MySQL，后台任务在事务提交后异步调用平台"
        "无关的 HTTP 通知网关。发送失败只更新 Outbox，不回滚采购业务。",
    )

    add_paragraph_before(error_heading, "9.7.1 查询通知记录", "Heading 3")
    add_endpoint_table_before(
        document,
        error_heading,
        method="GET",
        path="/api/v1/notifications",
        purpose="按 status、request_id 分页查询通知、失败原因和重试状态。",
    )
    add_paragraph_before(
        error_heading,
        "仅管理员可查询。status 支持 PENDING、SENT、FAILED；page_size 默认 20，最大 100。",
    )

    add_paragraph_before(error_heading, "9.7.2 触发到期通知发送", "Heading 3")
    add_endpoint_table_before(
        document,
        error_heading,
        method="POST",
        path="/api/v1/notifications/dispatch-due",
        purpose="管理员在开发或运维场景中立即处理一批到期通知。",
    )
    add_paragraph_before(
        error_heading,
        '请求体为 {"batch_size": 50}，范围 1 至 200。正式部署的定时任务可直接'
        "调用同一 Service，不必经过 HTTP 接口。",
    )

    add_paragraph_before(error_heading, "9.7.3 人工补发失败通知", "Heading 3")
    add_endpoint_table_before(
        document,
        error_heading,
        method="POST",
        path="/api/v1/notifications/{notification_id}/resend",
        purpose="复用原通知记录，将达到上限或需要人工处理的失败通知重新入队。",
    )
    add_paragraph_before(
        error_heading,
        '请求体包含 reason 和 action_token，例如 {"reason": "接收账号已恢复", '
        '"action_token": "<UUID>"}。只允许补发 FAILED 记录；操作写入采购操作日志，'
        "相同 action_token 不得重复执行。",
    )

    add_paragraph_before(error_heading, "9.7.4 发送与重试规则", "Heading 3")
    add_paragraph_before(
        error_heading,
        "后台任务使用 SELECT FOR UPDATE SKIP LOCKED 领取 PENDING 或已到达 "
        "next_retry_at 的 FAILED 记录，避免多个工作进程重复发送同一通知。",
    )
    add_paragraph_before(
        error_heading,
        "HTTP 网关请求使用 dedup_key 作为 Idempotency-Key。默认最多自动尝试 "
        "5 次，首次失败后等待 60 秒，随后指数退避，单次最长等待 24 小时；"
        "达到上限后 next_retry_at 置空，等待管理员补发。",
    )
    add_paragraph_before(
        error_heading,
        "未配置 NOTIFICATION_GATEWAY_URL 时按真实失败处理并记录原因，不会将通知"
        "误标记为已发送。后台单次处理入口为 python -m app.workers.notifications。",
    )

    document.save(output)
    return output


if __name__ == "__main__":
    print(update_interface_document())
